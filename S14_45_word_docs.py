import docx

# Access document object
d = docx.Document(
    r'C:\Users\benne\OneDrive\Documents\GitHub\Automate_the_boring_stuff\demo.docx')
# Access 2nd paragraph object in document
p = d.paragraphs[1]
# Print the text value of the 2nd paragraph
print(d.paragraphs[1].text)
# Print the 1st run object of the paragraph
print(p.runs[0])
print(p.runs[0].text)
# Print the bold/italic boolean of the run
print(p.runs[0].bold)
print(p.runs[3].italic)
# Assign new values to run object and save as docx
p.runs[3].underline = True
p.runs[3].text = 'italic and underline'
d.save('demo2.docx')
# Check style of object and change its value
print(p.style)
p.style = 'Title'
d.save('demo3.docx')

# Create a new document
d = docx.Document()
# Add paragraph objects
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph("This is another paragraph")
# Add a run to the first paragraph object
p = d.paragraphs[0]
p.add_run('This is a new run')
# Set the new run to bold
p.runs[1].bold = True
d.save('demo4.docx')


def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    return '\n'.join(full_text)


print(get_text('demo.docx'))
